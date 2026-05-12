"""
RedactAI — FastAPI Backend powered by Microsoft Presidio
Production-grade PII detection & redaction API
"""

import os
import json
import time
import uuid
import hashlib
from datetime import datetime, timezone
from typing import Optional
from fastapi import FastAPI, File, UploadFile, HTTPException, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel

# ---- Supabase Setup (with timeout to prevent startup hang) ----
SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://amklcfiyxeomdueeptyu.supabase.co")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImFta2xjZml5eGVvbWR1ZWVwdHl1Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3Nzg1MDQ2OTIsImV4cCI6MjA5NDA4MDY5Mn0.WNvPc9hrorOw_pMI2PS8pVPklfqwXCQH3kBJSwja6dk")
SUPABASE_AVAILABLE = False
supabase = None
try:
    from supabase import create_client
    import threading
    _sb_result = [False]
    def _check_sb():
        try:
            sb = create_client(SUPABASE_URL, SUPABASE_KEY)
            sb.table("redact_scans").select("id").limit(1).execute()
            _sb_result[0] = sb
        except:
            pass
    t = threading.Thread(target=_check_sb, daemon=True)
    t.start()
    t.join(timeout=5)  # Max 5 seconds for Supabase check
    if _sb_result[0]:
        supabase = _sb_result[0]
        SUPABASE_AVAILABLE = True
        print("[+] Supabase connected! Persistent history enabled.")
    else:
        print("[!] Supabase timed out or failed, falling back to in-memory history")
except Exception as e:
    print(f"[!] Supabase unavailable ({e}), falling back to in-memory history")

# ---- Presidio Setup ----
from presidio_analyzer import AnalyzerEngine, RecognizerRegistry
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

# ---- Piiranha Transformer Recognizer ----
from presidio_analyzer import EntityRecognizer, RecognizerResult

PIIRANHA_AVAILABLE = False
if os.environ.get("LOAD_PIIRANHA", "0") == "1":
    try:
        from transformers import pipeline as hf_pipeline
        PIIRANHA_AVAILABLE = True
    except ImportError:
        print("[!] transformers not installed, skipping Piiranha model")
else:
    print("[*] Piiranha model disabled (set LOAD_PIIRANHA=1 to enable)")

class PiiranhaRecognizer(EntityRecognizer):
    """Custom Presidio recognizer using the Piiranha PII model (DeBERTa-v3, 99.4% accuracy)"""

    PIIRANHA_TO_PRESIDIO = {
        "GIVENNAME": "PERSON",
        "SURNAME": "PERSON",
        "FIRSTNAME": "PERSON",
        "LASTNAME": "PERSON",
        "EMAIL": "EMAIL_ADDRESS",
        "PHONE": "PHONE_NUMBER",
        "PHONENUMBER": "PHONE_NUMBER",
        "CREDITCARD": "CREDIT_CARD",
        "CREDITCARDNUMBER": "CREDIT_CARD",
        "SOCIALNUM": "US_SSN",
        "SOCIALSECURITYNUMBER": "US_SSN",
        "DRIVERSLICENSE": "US_DRIVER_LICENSE",
        "DATEOFBIRTH": "DATE_TIME",
        "DOB": "DATE_TIME",
        "IDCARD": "ID_CARD",
        "TAXNUMBER": "TAX_ID",
        "STREETADDRESS": "LOCATION",
        "CITY": "LOCATION",
        "ZIPCODE": "LOCATION",
        "BUILDINGNUMBER": "LOCATION",
        "ACCOUNTNUMBER": "ACCOUNT_NUMBER",
        "USERNAME": "USERNAME",
        "PASSWORD": "PASSWORD",
    }

    def __init__(self):
        supported = list(set(self.PIIRANHA_TO_PRESIDIO.values()))
        super().__init__(
            supported_entities=supported,
            supported_language="en",
            name="PiiranhaRecognizer",
        )
        print("[*] Loading Piiranha PII transformer model...")
        self.pipe = hf_pipeline(
            "token-classification",
            model="iiiorg/piiranha-v1-detect-personal-information",
            aggregation_strategy="max",
            device=-1,  # CPU
        )
        print("[+] Piiranha model loaded!")

    def load(self):
        pass

    def analyze(self, text, entities=None, nlp_artifacts=None):
        results = []
        try:
            preds = self.pipe(text)
            for pred in preds:
                label = pred["entity_group"].upper().replace("-", "")
                presidio_type = self.PIIRANHA_TO_PRESIDIO.get(label, None)
                if presidio_type and (entities is None or presidio_type in entities):
                    results.append(
                        RecognizerResult(
                            entity_type=presidio_type,
                            start=pred["start"],
                            end=pred["end"],
                            score=round(float(pred["score"]), 3),
                        )
                    )
        except Exception as e:
            print(f"[!] Piiranha error: {e}")
        return results


# Initialize engines — try large model, fall back to small
print("[*] Loading NLP model & Presidio engines...")

# Try en_core_web_lg first (better NER), fall back to en_core_web_sm
for model_name in ["en_core_web_lg", "en_core_web_sm"]:
    try:
        nlp_config = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": model_name}],
        }
        nlp_engine = NlpEngineProvider(nlp_configuration=nlp_config).create_engine()
        print(f"[+] Using spaCy model: {model_name}")
        break
    except Exception as e:
        print(f"[!] {model_name} not available: {e}")
        continue

registry = RecognizerRegistry()
registry.load_predefined_recognizers(nlp_engine=nlp_engine)

# Add Piiranha transformer if available
if PIIRANHA_AVAILABLE:
    try:
        piiranha = PiiranhaRecognizer()
        registry.add_recognizer(piiranha)
        print("[+] Piiranha transformer recognizer added!")
    except Exception as e:
        print(f"[!] Could not load Piiranha model: {e}")
        print("[*] Continuing with spaCy-only detection")

# ---- GLiNER Zero-Shot NER (contextual understanding) ----
GLINER_AVAILABLE = False
if os.environ.get("LOAD_GLINER", "0") == "1":
    try:
        from gliner import GLiNER as GLiNERModel
        GLINER_AVAILABLE = True
    except ImportError:
        print("[!] gliner not installed, skipping zero-shot NER")
else:
    print("[*] GLiNER model disabled (set LOAD_GLINER=1 to enable)")

class GLiNERRecognizer(EntityRecognizer):
    """Zero-shot NER using GLiNER — understands context, no training needed.
    Detects entities based on natural language labels like 'person name', 'date'."""

    # Map GLiNER labels → Presidio entity types
    LABEL_MAP = {
        "person name": "PERSON",
        "full name": "PERSON",
        "date": "DATE_TIME",
        "monetary amount": "MONETARY_VALUE",
        "organization": "ORGANIZATION",
        "address": "LOCATION",
        "city": "LOCATION",
        "country": "LOCATION",
    }

    # Entity labels we ask GLiNER to find — written in natural language
    DETECT_LABELS = [
        "person name",
        "date",
        "monetary amount",
        "organization",
        "address",
    ]

    def __init__(self):
        supported = list(set(self.LABEL_MAP.values()))
        super().__init__(
            supported_entities=supported,
            supported_language="en",
            name="GLiNERRecognizer",
        )
        print("[*] Loading GLiNER zero-shot NER model...")
        self.model = GLiNERModel.from_pretrained("urchade/gliner_medium-v2.1")
        print("[+] GLiNER model loaded!")

    def load(self):
        pass

    def analyze(self, text, entities=None, nlp_artifacts=None):
        results = []
        try:
            preds = self.model.predict_entities(text, self.DETECT_LABELS, threshold=0.4)
            for pred in preds:
                label = pred["label"].lower()
                presidio_type = self.LABEL_MAP.get(label, None)
                if presidio_type and (entities is None or presidio_type in entities):
                    results.append(
                        RecognizerResult(
                            entity_type=presidio_type,
                            start=pred["start"],
                            end=pred["end"],
                            score=round(float(pred["score"]), 3),
                        )
                    )
        except Exception as e:
            print(f"[!] GLiNER error: {e}")
        return results

if GLINER_AVAILABLE:
    try:
        gliner_rec = GLiNERRecognizer()
        registry.add_recognizer(gliner_rec)
        print("[+] GLiNER zero-shot recognizer added!")
    except Exception as e:
        print(f"[!] Could not load GLiNER: {e}")

# ---- Custom Informal Date Recognizer ----
import re
from presidio_analyzer import Pattern, PatternRecognizer

MONTHS = r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"

informal_date_patterns = [
    # "9th march", "10th march", "1st january", "23rd april"
    Pattern("ordinal_month", rf"\b\d{{1,2}}(?:st|nd|rd|th)\s+{MONTHS}\b", 0.85),
    # "march 9th", "april 10th", "january 1st"
    Pattern("month_ordinal", rf"\b{MONTHS}\s+\d{{1,2}}(?:st|nd|rd|th)?\b", 0.85),
    # "march 2024", "april 2025"
    Pattern("month_year", rf"\b{MONTHS}\s+\d{{4}}\b", 0.80),
    # "9th march 2024"
    Pattern("ordinal_month_year", rf"\b\d{{1,2}}(?:st|nd|rd|th)\s+{MONTHS}\s+\d{{4}}\b", 0.90),
    # standalone months in context: "in march", "on april", "by december"
    Pattern("standalone_month", rf"\b(?:in|on|by|before|after|since|until|during)\s+{MONTHS}\b", 0.70),
]

date_recognizer = PatternRecognizer(
    supported_entity="DATE_TIME",
    name="InformalDateRecognizer",
    patterns=informal_date_patterns,
    supported_language="en",
)
registry.add_recognizer(date_recognizer)
print("[+] Informal date recognizer added!")

analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry, supported_languages=["en"])
anonymizer = AnonymizerEngine()
print("[+] Presidio engines ready!")

# ---- FastAPI App ----
app = FastAPI(
    title="RedactAI API",
    description="AI-powered PII detection & redaction API backed by Microsoft Presidio",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---- Storage (Supabase persistent + in-memory fallback) ----
scan_history_mem = []  # fallback only
api_keys = {
    "rda_live_sk_demo123": {"name": "Demo Key", "created": datetime.now().isoformat(), "active": True}
}

def save_scan(record):
    """Save a scan record to Supabase (or in-memory fallback)"""
    if SUPABASE_AVAILABLE:
        try:
            supabase.table("redact_scans").insert({
                "source": record["source"],
                "entity_count": record["entity_count"],
                "types": json.dumps(record["types"]),
                "processing_ms": int(record["processing_ms"]),
                "preview": record.get("preview", ""),
            }).execute()
            return
        except Exception as e:
            print(f"[!] Supabase insert failed: {e}")
    scan_history_mem.append(record)

# ---- Entity color/icon mapping for frontend ----
ENTITY_META = {
    "PERSON": {"icon": "👤", "color": "#f472b6", "cssClass": "name", "label": "Person Name"},
    "EMAIL_ADDRESS": {"icon": "📧", "color": "#74c0fc", "cssClass": "email", "label": "Email"},
    "PHONE_NUMBER": {"icon": "📱", "color": "#51cf66", "cssClass": "phone", "label": "Phone"},
    "CREDIT_CARD": {"icon": "💳", "color": "#ffd43b", "cssClass": "credit-card", "label": "Credit Card"},
    "US_SSN": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "SSN"},
    "US_PASSPORT": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Passport"},
    "US_DRIVER_LICENSE": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Driver License"},
    "IP_ADDRESS": {"icon": "🌐", "color": "#22d3ee", "cssClass": "ip", "label": "IP Address"},
    "DATE_TIME": {"icon": "📅", "color": "#a29bfe", "cssClass": "date", "label": "Date/Time"},
    "LOCATION": {"icon": "📍", "color": "#fdcb6e", "cssClass": "location", "label": "Location"},
    "NRP": {"icon": "🏛️", "color": "#dfe6e9", "cssClass": "other", "label": "Nationality/Religion"},
    "MEDICAL_LICENSE": {"icon": "🏥", "color": "#e17055", "cssClass": "other", "label": "Medical License"},
    "URL": {"icon": "🔗", "color": "#74c0fc", "cssClass": "other", "label": "URL"},
    "IBAN_CODE": {"icon": "🏦", "color": "#ffd43b", "cssClass": "credit-card", "label": "IBAN"},
    "CRYPTO": {"icon": "₿", "color": "#f9ca24", "cssClass": "other", "label": "Crypto Wallet"},
    "UK_NHS": {"icon": "🏥", "color": "#e17055", "cssClass": "gov-id", "label": "UK NHS Number"},
    "IN_AADHAAR": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Aadhaar"},
    "IN_PAN": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "PAN Card"},
    "ID_CARD": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "ID Card"},
    "TAX_ID": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Tax Number"},
    "ACCOUNT_NUMBER": {"icon": "🏦", "color": "#ffd43b", "cssClass": "credit-card", "label": "Account Number"},
    "USERNAME": {"icon": "👤", "color": "#a29bfe", "cssClass": "name", "label": "Username"},
    "PASSWORD": {"icon": "🔒", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Password"},
    "MONETARY_VALUE": {"icon": "💰", "color": "#ffd43b", "cssClass": "credit-card", "label": "Money/Amount"},
    "ORGANIZATION": {"icon": "🏢", "color": "#dfe6e9", "cssClass": "other", "label": "Organization"},
}

# ---- File Text Extraction ----
def extract_text_from_file(content: bytes, ext: str) -> str:
    """Extract text from various file formats"""
    import io
    
    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            print(f"[!] PDF extraction failed: {e}")
            return ""
    
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"[!] DOCX extraction failed: {e}")
            return ""
    
    elif ext in ("xlsx", "xls"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            text_parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        text_parts.append(" ".join(cells))
            wb.close()
            return "\n".join(text_parts)
        except Exception as e:
            print(f"[!] XLSX extraction failed: {e}")
            return ""
    
    elif ext == "csv":
        import csv
        text = content.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        return " ".join(" ".join(row) for row in reader)
    
    elif ext == "json":
        text = content.decode("utf-8", errors="ignore")
        try:
            data = json.loads(text)
            return json.dumps(data) if isinstance(data, (dict, list)) else text
        except json.JSONDecodeError:
            return text
    
    else:  # txt and fallback
        return content.decode("utf-8", errors="ignore")

# ---- Request/Response Models ----
class ScanRequest(BaseModel):
    text: str
    mode: str = "highlight"  # "highlight" or "redact"
    language: str = "en"
    entities: Optional[list] = None  # specific entities to detect, or None for all
    score_threshold: float = 0.35

class ScanResponse(BaseModel):
    original: str
    redacted: str
    entities: list
    entity_summary: dict
    count: int
    processing_ms: float

class BatchScanRequest(BaseModel):
    texts: list[str]
    mode: str = "redact"
    language: str = "en"

# ---- API Endpoints ----

@app.get("/api/health")
def health_check():
    return {"status": "healthy", "engine": "presidio", "version": "1.0.0"}


@app.post("/api/v1/scan", response_model=ScanResponse)
def scan_text(req: ScanRequest):
    """Scan text for PII and return detected entities + redacted text"""
    start = time.time()
    
    # Analyze with Presidio
    results = analyzer.analyze(
        text=req.text,
        language=req.language,
        entities=req.entities,
        score_threshold=req.score_threshold,
    )
    
    # Build entity list with metadata
    entities = []
    for r in sorted(results, key=lambda x: x.start):
        meta = ENTITY_META.get(r.entity_type, {"icon": "❓", "color": "#dfe6e9", "cssClass": "other", "label": r.entity_type})
        entities.append({
            "type": r.entity_type,
            "label": meta["label"],
            "text": req.text[r.start:r.end],
            "start": r.start,
            "end": r.end,
            "score": round(float(r.score), 3),
            "icon": meta["icon"],
            "color": meta["color"],
            "cssClass": meta["cssClass"],
        })
    
    # Anonymize/redact
    anonymized = anonymizer.anonymize(
        text=req.text,
        analyzer_results=results,
        operators={
            "DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"}),
            "PERSON": OperatorConfig("replace", {"new_value": "[NAME]"}),
            "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[EMAIL]"}),
            "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[PHONE]"}),
            "CREDIT_CARD": OperatorConfig("replace", {"new_value": "[CREDIT_CARD]"}),
            "US_SSN": OperatorConfig("replace", {"new_value": "[SSN]"}),
            "IP_ADDRESS": OperatorConfig("replace", {"new_value": "[IP_ADDRESS]"}),
            "DATE_TIME": OperatorConfig("replace", {"new_value": "[DATE]"}),
            "LOCATION": OperatorConfig("replace", {"new_value": "[LOCATION]"}),
            "URL": OperatorConfig("replace", {"new_value": "[URL]"}),
            "IN_AADHAAR": OperatorConfig("replace", {"new_value": "[AADHAAR]"}),
            "IN_PAN": OperatorConfig("replace", {"new_value": "[PAN]"}),
        }
    )
    
    # Build entity summary
    summary = {}
    for e in entities:
        t = e["label"]
        if t not in summary:
            summary[t] = {"count": 0, "icon": e["icon"], "cssClass": e["cssClass"]}
        summary[t]["count"] += 1
    
    elapsed_ms = round((time.time() - start) * 1000, 2)
    
    # Store in history (persistent via Supabase)
    save_scan({
        "id": str(uuid.uuid4())[:8],
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "source": "Text Input",
        "entity_count": len(entities),
        "types": list(summary.keys()),
        "processing_ms": elapsed_ms,
        "preview": req.text[:80] + ("..." if len(req.text) > 80 else ""),
    })
    
    return ScanResponse(
        original=req.text,
        redacted=anonymized.text,
        entities=entities,
        entity_summary=summary,
        count=len(entities),
        processing_ms=elapsed_ms,
    )


@app.post("/api/v1/scan/batch")
def scan_batch(req: BatchScanRequest):
    """Scan multiple texts at once"""
    results = []
    total_start = time.time()
    
    for text in req.texts:
        analysis = analyzer.analyze(text=text, language=req.language)
        anonymized = anonymizer.anonymize(text=text, analyzer_results=analysis)
        
        entities = []
        for r in analysis:
            meta = ENTITY_META.get(r.entity_type, {"icon": "❓", "label": r.entity_type})
            entities.append({
                "type": r.entity_type,
                "label": meta["label"],
                "text": text[r.start:r.end],
                "score": round(float(r.score), 3),
            })
        
        results.append({
            "original": text,
            "redacted": anonymized.text,
            "entity_count": len(entities),
            "entities": entities,
        })
    
    return {
        "results": results,
        "total_texts": len(req.texts),
        "total_entities": sum(r["entity_count"] for r in results),
        "processing_ms": round((time.time() - total_start) * 1000, 2),
    }


@app.post("/api/v1/scan/file")
async def scan_file(file: UploadFile = File(...)):
    """Upload and scan a file for PII — supports TXT, CSV, JSON, PDF, DOCX, XLSX"""
    if not file.filename:
        raise HTTPException(400, "No file provided")
    
    ext = file.filename.rsplit(".", 1)[-1].lower()
    supported = ("txt", "csv", "json", "pdf", "docx", "doc", "xlsx", "xls")
    if ext not in supported:
        raise HTTPException(400, f"Unsupported file type: .{ext}. Supported: {', '.join(supported)}")
    
    content = await file.read()
    start = time.time()
    
    # Extract text based on file type
    all_text = extract_text_from_file(content, ext)
    if not all_text or not all_text.strip():
        raise HTTPException(400, "Could not extract text from file")
    
    # Analyze
    results = analyzer.analyze(text=all_text, language="en")
    anonymized = anonymizer.anonymize(text=all_text, analyzer_results=results)
    
    entities = []
    for r in sorted(results, key=lambda x: x.start):
        meta = ENTITY_META.get(r.entity_type, {"icon": "❓", "label": r.entity_type})
        entities.append({
            "type": r.entity_type,
            "label": meta["label"],
            "text": all_text[r.start:r.end],
            "score": round(float(r.score), 3),
        })
    
    elapsed_ms = round((time.time() - start) * 1000, 2)
    
    # Store in history (persistent via Supabase)
    save_scan({
        "id": str(uuid.uuid4())[:8],
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "source": f"File: {file.filename}",
        "entity_count": len(entities),
        "types": list(set(e["label"] for e in entities)),
        "processing_ms": elapsed_ms,
        "preview": all_text[:80] + "...",
    })
    
    return {
        "filename": file.filename,
        "file_size": len(content),
        "redacted_text": anonymized.text,
        "entities": entities,
        "entity_count": len(entities),
        "processing_ms": elapsed_ms,
    }


@app.get("/api/v1/history")
def get_history(page: int = 1, per_page: int = 10):
    """Get scan history with pagination — reads from Supabase"""
    if SUPABASE_AVAILABLE:
        try:
            count_resp = supabase.table("redact_scans").select("id", count="exact").execute()
            total = count_resp.count or 0
            offset = (page - 1) * per_page
            data_resp = supabase.table("redact_scans") \
                .select("*") \
                .order("created_at", desc=True) \
                .range(offset, offset + per_page - 1) \
                .execute()
            items = []
            for row in data_resp.data:
                types_val = row.get("types", "[]")
                if isinstance(types_val, str):
                    try:
                        types_val = json.loads(types_val)
                    except Exception:
                        types_val = []
                items.append({
                    "id": str(row["id"])[:8],
                    "timestamp": row["created_at"],
                    "source": row.get("source", "Unknown"),
                    "entity_count": row.get("entity_count", 0),
                    "types": types_val,
                    "processing_ms": row.get("processing_ms", 0),
                    "preview": row.get("preview", ""),
                })
            return {
                "items": items,
                "total": total,
                "page": page,
                "pages": max(1, (total + per_page - 1) // per_page),
            }
        except Exception as e:
            print(f"[!] Supabase history read failed: {e}")
    # Fallback to in-memory
    total = len(scan_history_mem)
    start = (page - 1) * per_page
    items = list(reversed(scan_history_mem))[start:start + per_page]
    return {
        "items": items,
        "total": total,
        "page": page,
        "pages": max(1, (total + per_page - 1) // per_page),
    }


@app.get("/api/v1/stats")
def get_stats():
    """Get overview statistics — reads from Supabase"""
    if SUPABASE_AVAILABLE:
        try:
            count_resp = supabase.table("redact_scans").select("id", count="exact").execute()
            total_scans = count_resp.count or 0
            all_resp = supabase.table("redact_scans").select("entity_count,processing_ms,types").execute()
            rows = all_resp.data or []
            total_entities = sum(r.get("entity_count", 0) for r in rows)
            avg_ms = round(sum(r.get("processing_ms", 0) for r in rows) / max(1, total_scans), 2)
            type_counts = {}
            for r in rows:
                types_val = r.get("types", "[]")
                if isinstance(types_val, str):
                    try:
                        types_val = json.loads(types_val)
                    except Exception:
                        types_val = []
                for t in types_val:
                    type_counts[t] = type_counts.get(t, 0) + 1
            return {
                "total_scans": total_scans,
                "total_entities": total_entities,
                "avg_response_ms": avg_ms,
                "entity_type_breakdown": type_counts,
            }
        except Exception as e:
            print(f"[!] Supabase stats read failed: {e}")
    # Fallback to in-memory
    total_scans = len(scan_history_mem)
    total_entities = sum(h["entity_count"] for h in scan_history_mem)
    avg_ms = round(sum(h["processing_ms"] for h in scan_history_mem) / max(1, total_scans), 2)
    type_counts = {}
    for h in scan_history_mem:
        for t in h.get("types", []):
            type_counts[t] = type_counts.get(t, 0) + 1
    return {
        "total_scans": total_scans,
        "total_entities": total_entities,
        "avg_response_ms": avg_ms,
        "entity_type_breakdown": type_counts,
    }


@app.get("/api/v1/supported-entities")
def get_supported_entities():
    """List all PII entity types the engine can detect"""
    supported = analyzer.get_supported_entities()
    entities = []
    for entity_type in sorted(supported):
        meta = ENTITY_META.get(entity_type, {"icon": "❓", "color": "#dfe6e9", "label": entity_type})
        entities.append({
            "type": entity_type,
            "label": meta["label"],
            "icon": meta["icon"],
            "color": meta["color"],
        })
    return {"entities": entities, "count": len(entities)}


class CustomDetectorRequest(BaseModel):
    name: str
    entity_type: str
    regex: str
    score: float = 0.8


@app.post("/api/v1/custom-detector")
def add_custom_detector(req: CustomDetectorRequest):
    """Register a custom regex-based PII detector at runtime"""
    import re
    
    # Validate regex
    try:
        re.compile(req.regex)
    except re.error as e:
        raise HTTPException(400, f"Invalid regex: {e}")
    
    # Use Presidio's PatternRecognizer API
    from presidio_analyzer import Pattern, PatternRecognizer
    
    pattern = Pattern(
        name=req.name,
        regex=req.regex,
        score=req.score,
    )
    
    recognizer = PatternRecognizer(
        supported_entity=req.entity_type,
        name=f"custom_{req.name.lower().replace(' ', '_')}",
        patterns=[pattern],
    )
    
    # Add to the live registry
    analyzer.registry.add_recognizer(recognizer)
    
    # Also add to entity meta for frontend display
    ENTITY_META[req.entity_type] = {
        "icon": "🔧",
        "color": "#b8e994",
        "cssClass": "other",
        "label": req.name,
    }
    
    return {
        "status": "ok",
        "message": f"Custom detector '{req.name}' registered for entity '{req.entity_type}'",
        "entity_type": req.entity_type,
        "pattern": req.regex,
    }


@app.get("/api/v1/custom-detectors")
def list_custom_detectors():
    """List all custom detectors currently registered"""
    custom = []
    for rec in analyzer.registry.recognizers:
        if hasattr(rec, 'name') and rec.name and rec.name.startswith('custom_'):
            patterns = []
            if hasattr(rec, 'patterns'):
                patterns = [{"name": p.name, "regex": p.regex, "score": p.score} for p in rec.patterns]
            custom.append({
                "name": rec.name,
                "entity_type": rec.supported_entities[0] if rec.supported_entities else "UNKNOWN",
                "patterns": patterns,
            })
    return {"detectors": custom, "count": len(custom)}


@app.get("/api/v1/export")
def export_history(format: str = "csv"):
    """Export scan history as CSV or JSON — for compliance/audit"""
    import io
    
    # Fetch all history
    items = []
    if SUPABASE_AVAILABLE:
        try:
            resp = supabase.table("redact_scans") \
                .select("*") \
                .order("created_at", desc=True) \
                .limit(1000) \
                .execute()
            for row in resp.data:
                types_val = row.get("types", "[]")
                if isinstance(types_val, str):
                    try:
                        types_val = json.loads(types_val)
                    except Exception:
                        types_val = []
                items.append({
                    "id": str(row["id"])[:8],
                    "timestamp": row["created_at"],
                    "source": row.get("source", ""),
                    "entity_count": row.get("entity_count", 0),
                    "types": ", ".join(types_val) if types_val else "",
                    "processing_ms": row.get("processing_ms", 0),
                    "preview": row.get("preview", ""),
                })
        except Exception as e:
            print(f"[!] Export from Supabase failed: {e}")
    else:
        for h in reversed(scan_history_mem):
            items.append({
                "id": h.get("id", ""),
                "timestamp": h.get("timestamp", ""),
                "source": h.get("source", ""),
                "entity_count": h.get("entity_count", 0),
                "types": ", ".join(h.get("types", [])),
                "processing_ms": h.get("processing_ms", 0),
                "preview": h.get("preview", ""),
            })
    
    if format == "json":
        return JSONResponse(content={"export": items, "total": len(items), "exported_at": datetime.now(timezone.utc).isoformat()})
    
    # CSV format
    import csv
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=["id", "timestamp", "source", "entity_count", "types", "processing_ms", "preview"])
    writer.writeheader()
    writer.writerows(items)
    
    from fastapi.responses import StreamingResponse
    csv_content = output.getvalue()
    return StreamingResponse(
        iter([csv_content]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=redactai_audit_log_{datetime.now().strftime('%Y%m%d')}.csv"}
    )


# =============================================
# SHADOW AI / WEBSITE PRIVACY SCANNER
# Dual-engine: Jina Reader API (JS-rendered text)
# + requests/BS4 (raw HTML tracker analysis).
# Production-grade — works on any cloud platform.
# Inspired by The Markup's Blacklight scanner.
# =============================================

# Known tracker signatures — domain patterns and their categories
TRACKER_SIGNATURES = {
    # Analytics
    "google-analytics.com": {"name": "Google Analytics", "category": "analytics", "risk": "medium"},
    "googletagmanager.com": {"name": "Google Tag Manager", "category": "analytics", "risk": "medium"},
    "analytics.google.com": {"name": "Google Analytics", "category": "analytics", "risk": "medium"},
    "gtag/js": {"name": "Google Global Site Tag", "category": "analytics", "risk": "medium"},
    "plausible.io": {"name": "Plausible Analytics", "category": "analytics", "risk": "low"},
    "umami.is": {"name": "Umami Analytics", "category": "analytics", "risk": "low"},
    "matomo": {"name": "Matomo Analytics", "category": "analytics", "risk": "low"},
    "mixpanel.com": {"name": "Mixpanel", "category": "analytics", "risk": "high"},
    "segment.com": {"name": "Segment", "category": "analytics", "risk": "high"},
    "amplitude.com": {"name": "Amplitude", "category": "analytics", "risk": "high"},
    "heap-analytics": {"name": "Heap Analytics", "category": "analytics", "risk": "high"},
    "heapanalytics.com": {"name": "Heap Analytics", "category": "analytics", "risk": "high"},
    "clarity.ms": {"name": "Microsoft Clarity", "category": "session_recording", "risk": "high"},
    # Advertising / Retargeting
    "facebook.net": {"name": "Meta Pixel (Facebook)", "category": "advertising", "risk": "high"},
    "facebook.com/tr": {"name": "Meta Pixel Tracking", "category": "advertising", "risk": "high"},
    "fbevents.js": {"name": "Meta Pixel Events", "category": "advertising", "risk": "high"},
    "connect.facebook": {"name": "Facebook Connect", "category": "advertising", "risk": "high"},
    "doubleclick.net": {"name": "Google Ads (DoubleClick)", "category": "advertising", "risk": "high"},
    "googlesyndication.com": {"name": "Google AdSense", "category": "advertising", "risk": "high"},
    "googleadservices.com": {"name": "Google Ads Conversion", "category": "advertising", "risk": "high"},
    "ads-twitter.com": {"name": "X (Twitter) Ads", "category": "advertising", "risk": "high"},
    "analytics.tiktok.com": {"name": "TikTok Pixel", "category": "advertising", "risk": "high"},
    "snap.licdn.com": {"name": "LinkedIn Insight Tag", "category": "advertising", "risk": "high"},
    "px.ads.linkedin.com": {"name": "LinkedIn Ads Pixel", "category": "advertising", "risk": "high"},
    "ads.reddit.com": {"name": "Reddit Pixel", "category": "advertising", "risk": "medium"},
    "static.criteo.net": {"name": "Criteo Retargeting", "category": "advertising", "risk": "high"},
    "bat.bing.com": {"name": "Microsoft Ads UET", "category": "advertising", "risk": "medium"},
    # Session Recording
    "hotjar.com": {"name": "Hotjar", "category": "session_recording", "risk": "high"},
    "fullstory.com": {"name": "FullStory", "category": "session_recording", "risk": "high"},
    "mouseflow.com": {"name": "Mouseflow", "category": "session_recording", "risk": "high"},
    "smartlook.com": {"name": "Smartlook", "category": "session_recording", "risk": "high"},
    "logrocket.com": {"name": "LogRocket", "category": "session_recording", "risk": "high"},
    "inspectlet.com": {"name": "Inspectlet", "category": "session_recording", "risk": "high"},
    # Customer Data Platforms
    "intercom.io": {"name": "Intercom", "category": "cdp", "risk": "medium"},
    "drift.com": {"name": "Drift Chat", "category": "cdp", "risk": "medium"},
    "hubspot.com": {"name": "HubSpot", "category": "cdp", "risk": "medium"},
    "hs-scripts.com": {"name": "HubSpot Scripts", "category": "cdp", "risk": "medium"},
    "crisp.chat": {"name": "Crisp Chat", "category": "cdp", "risk": "medium"},
    "tawk.to": {"name": "Tawk.to Chat", "category": "cdp", "risk": "low"},
    "zendesk.com": {"name": "Zendesk", "category": "cdp", "risk": "medium"},
    # Fingerprinting
    "fingerprintjs": {"name": "FingerprintJS", "category": "fingerprinting", "risk": "high"},
    "fpjs.io": {"name": "Fingerprint Pro", "category": "fingerprinting", "risk": "high"},
}

# AI / LLM endpoint patterns
AI_ENDPOINT_PATTERNS = [
    "api.openai.com", "api.anthropic.com", "api.fireworks.ai",
    "api.together.xyz", "api.replicate.com", "api.groq.com",
    "generativelanguage.googleapis.com", "api.cohere.ai",
    "api-inference.huggingface.co", "api.mistral.ai",
    "chatgpt", "gpt-4", "gpt-3", "claude", "gemini",
    "sk-proj-", "sk-ant-", "sk_live_", "fw_",  # API key patterns
]

# ---- BLACKLIGHT-GRADE ADVANCED DETECTION PATTERNS ----
# Ported from The Markup's Blacklight methodology:
# https://themarkup.org/blacklight/2020/09/22/how-we-built-a-real-time-privacy-inspector

# Canvas fingerprinting — JS API calls that uniquely identify browsers
# (Blacklight's canvas_fingerprinters test)
CANVAS_FINGERPRINT_PATTERNS = [
    "toDataURL",            # HTMLCanvasElement.toDataURL() — exports canvas as image
    "getImageData",         # CanvasRenderingContext2D.getImageData() — reads pixel data
    "measureText",          # Used with specific fonts to detect installed fonts
    "isPointInPath",        # Geometry-based fingerprinting
    "isPointInStroke",
    "canvas.toBlob",        # Another canvas export method
    "OffscreenCanvas",      # Off-screen canvas (stealthier fingerprinting)
    "WebGLRenderingContext", # WebGL fingerprinting
    "WEBGL_debug_renderer_info", # GPU fingerprinting via WebGL
    "getExtension",         # WebGL extensions for fingerprinting
]

# Key logging — scripts that capture keystrokes before form submission
# (Blacklight's key_logging test)
KEYLOGGING_PATTERNS = [
    "addEventListener('keydown'",
    'addEventListener("keydown"',
    "addEventListener('keypress'",
    'addEventListener("keypress"',
    "addEventListener('keyup'",
    'addEventListener("keyup"',
    "addEventListener('input'",
    'addEventListener("input"',
    "onkeydown",
    "onkeypress",
    "onkeyup",
    "document.onkeydown",
    "document.onkeypress",
    "inputMode",
    "event.key",
    "event.keyCode",
    "event.charCode",
    "event.which",
]

# Session recorder deep patterns — scripts that record mouse/scroll/clicks
# (Blacklight's session_recorders test)
SESSION_RECORDER_PATTERNS = [
    # Mouse tracking
    "addEventListener('mousemove'",
    'addEventListener("mousemove"',
    "addEventListener('mousedown'",
    "addEventListener('mouseup'",
    "addEventListener('click'",
    "addEventListener('scroll'",
    "addEventListener('touchstart'",
    "addEventListener('touchmove'",
    # Known session recorder libraries
    "rrweb",                    # Open-source session recorder
    "rrwebPlayer",
    "__rrweb",
    "sessionstack.com",
    "decibelinsight.net",
    "quantummetric.com",
    "contentsquare.com",
    "glassbox.com",
    "clicktale.net",
    "crazyegg.com",
    "Lucky Orange",
    "luckyorange.com",
    # DOM mutation observation (used by recorders)
    "MutationObserver",
    "IntersectionObserver",
]

# Facebook Pixel deep event patterns (Blacklight's fb_pixel_events test)
FB_PIXEL_EVENTS = [
    "fbq('track'",
    'fbq("track"',
    "fbq('init'",
    'fbq("init"',
    "fbq('trackCustom'",
    "_fbq",
    "facebook.com/tr?",
    "PageView",            # FB standard events
    "ViewContent",
    "AddToCart",
    "Purchase",
    "CompleteRegistration",
    "Lead",
    "InitiateCheckout",
]

# Google Analytics deep event patterns (Blacklight's google_analytics_events test)
GA_EVENT_PATTERNS = [
    "gtag('event'",
    'gtag("event"',
    "gtag('config'",
    'gtag("config"',
    "ga('send'",
    'ga("send"',
    "ga('create'",
    "_gaq.push",
    "__gaTracker",
    "GoogleAnalyticsObject",
    "analytics.js",
    "measurement_id",
    "send_page_view",
    "page_view",
    "enhanced_conversions",
    "user_id",              # User ID tracking (high privacy risk)
    "client_id",
]

# Known third-party tracking domains (expanded from Blacklight + disconnect.me lists)
TRACKING_DOMAINS = [
    # Data brokers / ad networks
    "adnxs.com", "adsrvr.org", "casalemedia.com", "contextweb.com",
    "demdex.net", "dotomi.com", "exponential.com", "eyereturn.com",
    "indexexchange.com", "liadm.com", "mathtag.com", "mookie1.com",
    "openx.net", "pubmatic.com", "rlcdn.com", "rubiconproject.com",
    "scorecardresearch.com", "serving-sys.com", "sharethrough.com",
    "simpli.fi", "sitescout.com", "smartadserver.com", "taboola.com",
    "outbrain.com", "tapad.com", "turn.com", "quantserve.com",
    # Data management platforms
    "bluekai.com", "bombora.com", "demandbase.com", "everesttech.net",
    "krxd.net", "moatads.com", "narrative.io", "oracle.com/cx",
    # Social tracking
    "platform.twitter.com", "platform.linkedin.com", "connect.facebook.net",
    "platform.instagram.com", "apis.google.com/js/platform",
]

# PII-collecting input field patterns
PII_INPUT_PATTERNS = {
    "name": ["name", "fullname", "full_name", "firstname", "lastname", "first_name", "last_name", "your-name"],
    "email": ["email", "e-mail", "mail", "emailaddress", "email_address", "your-email"],
    "phone": ["phone", "tel", "telephone", "mobile", "cell", "phonenumber", "phone_number"],
    "address": ["address", "street", "city", "state", "zip", "zipcode", "postal", "country"],
    "dob": ["dob", "birthday", "birthdate", "date_of_birth", "dateofbirth"],
    "ssn": ["ssn", "social_security", "socialsecurity", "national_id", "nationalid"],
    "card": ["card", "credit_card", "creditcard", "cardnumber", "card_number", "cvv", "cvc", "expiry"],
    "password": ["password", "passwd", "pass", "secret"],
    "aadhaar": ["aadhaar", "aadhar", "uid_number"],
    "pan": ["pan_number", "pan_card", "pancard"],
}


class URLScanRequest(BaseModel):
    url: str
    email: Optional[str] = None


@app.post("/api/v1/scan/url")
async def scan_url(req: URLScanRequest):
    """
    Shadow AI / Website Privacy Scanner — Production-grade.
    Dual-engine approach:
      1. Jina Reader API (r.jina.ai) — free, cloud-hosted, handles JS/SPAs,
         returns clean text from ANY website. No API key needed.
      2. requests + BeautifulSoup — raw HTML analysis for trackers,
         forms, scripts, pixels, compliance checks.
    Then: Presidio NLP engine scans extracted text for PII.
    Works identically on local, HuggingFace, Vercel, any cloud.
    """
    import requests as http_requests
    from urllib.parse import urlparse
    from bs4 import BeautifulSoup

    url = req.url.strip()
    if not url.startswith("http"):
        url = "https://" + url

    parsed = urlparse(url)
    base_domain = parsed.netloc.lower()

    start_time = time.time()

    # ---- ENGINE 1: Raw HTML fetch (for tracker/form/script analysis) ----
    browser_headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
    }
    try:
        resp = http_requests.get(url, headers=browser_headers, timeout=20, allow_redirects=True, verify=True)
        html = resp.text
        final_url = str(resp.url)
        status_code = resp.status_code
        is_https = final_url.startswith("https://")
        response_headers = dict(resp.headers)
    except http_requests.exceptions.SSLError:
        try:
            resp = http_requests.get(url, headers=browser_headers, timeout=20, allow_redirects=True, verify=False)
            html = resp.text
            final_url = str(resp.url)
            status_code = resp.status_code
            is_https = False
            response_headers = dict(resp.headers)
        except Exception as e:
            raise HTTPException(400, f"Could not fetch URL: {str(e)}")
    except Exception as e:
        raise HTTPException(400, f"Could not fetch URL: {str(e)}")

    if status_code >= 400:
        raise HTTPException(400, f"URL returned HTTP {status_code}")

    # ---- ENGINE 2: Jina Reader API (deep JS-rendered text extraction) ----
    # Free, no API key, handles React/Vue/Angular/SPAs, returns clean text.
    # Falls back to BS4 text extraction if Jina is unreachable.
    jina_text = ""
    jina_used = False
    try:
        jina_url = f"https://r.jina.ai/{url}"
        jina_resp = http_requests.get(
            jina_url,
            headers={"Accept": "text/plain", "X-Return-Format": "text"},
            timeout=30,
        )
        if jina_resp.ok and len(jina_resp.text) > 100:
            jina_text = jina_resp.text
            jina_used = True
    except Exception:
        pass  # Fallback to BS4

    # ---- PARSE HTML ----
    soup_full = BeautifulSoup(html, "html.parser")

    # BS4 text extraction (fallback / supplement)
    soup_text = BeautifulSoup(html, "html.parser")
    for tag in soup_text(["script", "style", "noscript", "svg", "path"]):
        tag.decompose()
    bs4_text = soup_text.get_text(separator=" ", strip=True)

    # Use Jina text (deeper, JS-rendered) when available, otherwise BS4
    visible_text = jina_text if jina_used else bs4_text

    
    # ---- 3. DETECT TRACKERS ----
    trackers_found = []
    tracker_categories = {}
    all_scripts = soup_full.find_all("script", src=True)
    all_links = soup_full.find_all("link", href=True)
    all_imgs = soup_full.find_all("img", src=True)
    inline_scripts = soup_full.find_all("script", src=False)
    inline_script_text = " ".join([s.string or "" for s in inline_scripts])
    
    # Check all external resources
    all_src_urls = []
    for s in all_scripts:
        all_src_urls.append(s.get("src", ""))
    for l in all_links:
        all_src_urls.append(l.get("href", ""))
    for img in all_imgs:
        all_src_urls.append(img.get("src", ""))
    
    # Also check inline scripts
    full_check_text = " ".join(all_src_urls) + " " + inline_script_text
    
    seen_trackers = set()
    for signature, info in TRACKER_SIGNATURES.items():
        if signature.lower() in full_check_text.lower():
            if info["name"] not in seen_trackers:
                seen_trackers.add(info["name"])
                trackers_found.append({
                    "name": info["name"],
                    "category": info["category"],
                    "risk": info["risk"],
                    "signature": signature,
                })
                cat = info["category"]
                tracker_categories[cat] = tracker_categories.get(cat, 0) + 1
    
    # ---- 4. DETECT TRACKING PIXELS (1x1 images) ----
    tracking_pixels = []
    for img in all_imgs:
        src = img.get("src", "")
        width = img.get("width", "")
        height = img.get("height", "")
        style = img.get("style", "")
        is_pixel = False
        if (width == "1" and height == "1") or (width == "0" and height == "0"):
            is_pixel = True
        if "display:none" in style or "visibility:hidden" in style:
            is_pixel = True
        if is_pixel and src:
            tracking_pixels.append({"src": src[:200], "hidden": True})
    
    # ---- 5. DETECT DATA COLLECTION FORMS ----
    forms_found = []
    all_forms = soup_full.find_all("form")
    all_inputs = soup_full.find_all("input")
    
    pii_inputs_found = []
    for inp in all_inputs:
        input_name = (inp.get("name", "") or "").lower()
        input_type = (inp.get("type", "") or "").lower()
        input_id = (inp.get("id", "") or "").lower()
        input_placeholder = (inp.get("placeholder", "") or "").lower()
        check_str = f"{input_name} {input_type} {input_id} {input_placeholder}"
        
        for pii_type, patterns in PII_INPUT_PATTERNS.items():
            for pattern in patterns:
                if pattern in check_str:
                    pii_inputs_found.append({
                        "type": pii_type,
                        "field_name": input_name or input_id or input_placeholder[:40],
                        "input_type": input_type,
                    })
                    break
    
    # Deduplicate
    seen_inputs = set()
    unique_pii_inputs = []
    for inp in pii_inputs_found:
        key = f"{inp['type']}:{inp['field_name']}"
        if key not in seen_inputs:
            seen_inputs.add(key)
            unique_pii_inputs.append(inp)
    
    # ---- 6. DETECT AI/LLM ENDPOINTS ----
    ai_endpoints_found = []
    for pattern in AI_ENDPOINT_PATTERNS:
        if pattern.lower() in full_check_text.lower():
            # Determine if it's an API key leak vs endpoint reference
            is_key_leak = pattern.startswith("sk-") or pattern.startswith("fw_")
            ai_endpoints_found.append({
                "pattern": pattern,
                "type": "api_key_leak" if is_key_leak else "ai_endpoint",
                "risk": "critical" if is_key_leak else "high",
            })
    
    # ---- 7. BLACKLIGHT-GRADE: Canvas Fingerprinting Detection ----
    canvas_fp_signals = []
    for pattern in CANVAS_FINGERPRINT_PATTERNS:
        if pattern in inline_script_text:
            canvas_fp_signals.append(pattern)
    canvas_fingerprinting = len(canvas_fp_signals) >= 2  # Need 2+ signals to confirm
    
    # ---- 8. BLACKLIGHT-GRADE: Key Logging Detection ----
    keylog_signals = []
    for pattern in KEYLOGGING_PATTERNS:
        if pattern in inline_script_text or pattern in full_check_text:
            keylog_signals.append(pattern)
    key_logging_detected = len(keylog_signals) >= 2
    
    # ---- 9. BLACKLIGHT-GRADE: Session Recorder Deep Detection ----
    session_rec_signals = []
    for pattern in SESSION_RECORDER_PATTERNS:
        if pattern.lower() in full_check_text.lower() or pattern.lower() in inline_script_text.lower():
            session_rec_signals.append(pattern)
    session_recording_detected = len(session_rec_signals) >= 2
    
    # ---- 10. BLACKLIGHT-GRADE: Facebook Pixel Events ----
    fb_pixel_events = []
    for pattern in FB_PIXEL_EVENTS:
        if pattern in inline_script_text or pattern in full_check_text:
            fb_pixel_events.append(pattern)
    fb_pixel_detected = len(fb_pixel_events) > 0
    
    # ---- 11. BLACKLIGHT-GRADE: Google Analytics Events ----
    ga_events = []
    for pattern in GA_EVENT_PATTERNS:
        if pattern in inline_script_text or pattern in full_check_text:
            ga_events.append(pattern)
    ga_detected = len(ga_events) > 0
    
    # ---- 12. Third-party tracking domains (Disconnect.me list) ----
    third_party_domains_found = []
    for domain in TRACKING_DOMAINS:
        if domain.lower() in full_check_text.lower():
            third_party_domains_found.append(domain)
    
    # ---- 13. COMPLIANCE CHECKS ----
    page_lower = html.lower()
    
    has_privacy_policy = any(kw in page_lower for kw in [
        "privacy policy", "privacy-policy", "privacypolicy",
        "/privacy", "data protection", "datenschutz"
    ])
    
    has_cookie_consent = any(kw in page_lower for kw in [
        "cookie consent", "cookie-consent", "cookieconsent",
        "cookie banner", "cookie-banner", "cookie policy",
        "accept cookies", "cookie notice", "gdpr",
        "cookie-law", "cookie_consent", "onetrust",
        "cookiebot", "osano", "termly"
    ])
    
    has_terms = any(kw in page_lower for kw in [
        "terms of service", "terms-of-service", "terms and conditions",
        "terms-and-conditions", "/terms", "/tos"
    ])
    
    # Security headers — detailed analysis
    security_headers = {
        "content-security-policy": "Content-Security-Policy" in response_headers,
        "x-frame-options": "X-Frame-Options" in response_headers,
        "strict-transport-security": "Strict-Transport-Security" in response_headers,
        "x-content-type-options": "X-Content-Type-Options" in response_headers,
        "x-xss-protection": "X-XSS-Protection" in response_headers,
        "referrer-policy": "Referrer-Policy" in response_headers,
        "permissions-policy": "Permissions-Policy" in response_headers,
    }
    sec_header_score = sum(1 for v in security_headers.values() if v)
    sec_header_grade = "A" if sec_header_score >= 6 else "B" if sec_header_score >= 4 else "C" if sec_header_score >= 2 else "F"
    
    # ---- 14. DEEP PRIVACY POLICY SCAN ----
    # DPDP compliance signals are usually on /privacy or /terms pages, not the homepage.
    # To avoid false negatives, we also fetch the privacy policy page.
    privacy_page_text = ""
    privacy_url_found = None
    try:
        # Find privacy policy link from homepage
        for link in soup_full.find_all("a", href=True):
            href = link.get("href", "").lower()
            link_text = (link.get_text() or "").lower()
            if any(kw in href for kw in ["/privacy", "privacy-policy", "privacypolicy", "data-protection"]) or \
               any(kw in link_text for kw in ["privacy policy", "privacy notice", "data protection"]):
                privacy_href = link.get("href", "")
                # Resolve relative URLs
                if privacy_href.startswith("/"):
                    from urllib.parse import urlparse
                    parsed = urlparse(final_url)
                    privacy_url_found = f"{parsed.scheme}://{parsed.netloc}{privacy_href}"
                elif privacy_href.startswith("http"):
                    privacy_url_found = privacy_href
                break
        
        # Fetch privacy policy page
        if privacy_url_found:
            pp_resp = http_requests.get(privacy_url_found, headers={"User-Agent": "Mozilla/5.0 RedactAI-Scanner/2.0"}, timeout=10)
            if pp_resp.ok:
                pp_soup = BeautifulSoup(pp_resp.text, "html.parser")
                for tag in pp_soup(["script", "style", "noscript"]):
                    tag.decompose()
                privacy_page_text = pp_soup.get_text(separator=" ", strip=True).lower()
    except Exception as e:
        print(f"[!] Privacy policy page fetch failed: {e}")
    
    # Combine homepage + privacy page for DPDP analysis
    combined_compliance_text = page_lower + " " + privacy_page_text
    
    # ---- 15. DPDP ACT 2023 (India) COMPLIANCE CHECKS ----
    # Based on the Digital Personal Data Protection Act, 2023
    # Now checks BOTH homepage AND privacy policy page for accuracy
    dpdp_checks = {}
    
    # Consent mechanism — DPDP requires free, specific, informed, unconditional consent
    dpdp_checks["consent_mechanism"] = any(kw in combined_compliance_text for kw in [
        "i agree", "i consent", "accept cookies", "cookie consent",
        "by continuing", "by using this", "consent to",
        "opt-in", "opt in", "accept all", "reject all",
        "manage preferences", "cookie preferences", "cookie settings",
        "onetrust", "cookiebot", "osano", "termly", "truendo",
        "consent management", "lawful basis", "legal basis",
    ])
    
    # Privacy notice — DPDP Section 5: must inform purpose of data collection
    dpdp_checks["privacy_notice"] = has_privacy_policy
    
    # Grievance officer / DPO contact — DPDP Section 8(7)
    dpdp_checks["grievance_officer"] = any(kw in combined_compliance_text for kw in [
        "grievance officer", "grievance redressal", "data protection officer",
        "dpo@", "grievance@", "privacy@", "nodal officer",
        "grievance.officer", "data-protection-officer",
        "grievance mechanism", "redressal mechanism",
    ])
    
    # Data retention / deletion policy — DPDP Section 8(6)
    dpdp_checks["data_retention_policy"] = any(kw in combined_compliance_text for kw in [
        "data retention", "retention policy", "data deletion",
        "erase your data", "delete your data", "right to erasure",
        "right to be forgotten", "data erasure", "retain your",
        "retention period", "how long we keep", "how long we store",
        "stored for a period", "deleted after", "erasure of data",
    ])
    
    # Children's data protection — DPDP Section 9
    dpdp_checks["children_protection"] = any(kw in combined_compliance_text for kw in [
        "children", "child", "minor", "parental consent",
        "under 18", "under 13", "coppa", "age verification",
        "verifiable parental", "age gate", "minors",
    ])
    
    # Consent withdrawal mechanism — DPDP Section 6(4)
    dpdp_checks["consent_withdrawal"] = any(kw in combined_compliance_text for kw in [
        "withdraw consent", "revoke consent", "opt out", "opt-out",
        "unsubscribe", "manage consent", "withdraw your consent",
        "right to withdraw", "change your preferences",
        "modify your consent", "update your preferences",
    ])
    
    # Data breach notification reference — DPDP Section 8(5)
    dpdp_checks["breach_notification"] = any(kw in combined_compliance_text for kw in [
        "data breach", "breach notification", "security incident",
        "notify the board", "data protection board",
        "security breach", "breach of data", "unauthorized access",
        "incident response", "notify you of",
    ])
    
    dpdp_score = sum(1 for v in dpdp_checks.values() if v)
    dpdp_grade = "A" if dpdp_score >= 6 else "B" if dpdp_score >= 4 else "C" if dpdp_score >= 2 else "F"
    
    # ---- 15. COOKIE DEEP ANALYSIS (from Set-Cookie headers) ----
    cookie_analysis = []
    set_cookie_headers = response_headers.get("Set-Cookie", "") or response_headers.get("set-cookie", "")
    if isinstance(set_cookie_headers, str):
        set_cookie_headers = [set_cookie_headers] if set_cookie_headers else []
    
    for cookie_str in set_cookie_headers:
        if not cookie_str.strip():
            continue
        parts = cookie_str.split(";")
        name_val = parts[0].split("=", 1)
        cookie_name = name_val[0].strip() if name_val else "unknown"
        cookie_flags = cookie_str.lower()
        
        cookie_info = {
            "name": cookie_name[:40],
            "httponly": "httponly" in cookie_flags,
            "secure": "secure" in cookie_flags,
            "samesite": "samesite=strict" in cookie_flags or "samesite=lax" in cookie_flags,
            "third_party": base_domain not in cookie_str.lower(),
        }
        # Duration analysis
        if "max-age=" in cookie_flags:
            try:
                age = int(cookie_flags.split("max-age=")[1].split(";")[0].strip())
                cookie_info["duration_days"] = round(age / 86400, 1)
                cookie_info["persistent"] = age > 86400  # > 1 day
            except:
                cookie_info["persistent"] = True
        elif "expires=" in cookie_flags:
            cookie_info["persistent"] = True
        else:
            cookie_info["persistent"] = False  # Session cookie
        
        cookie_analysis.append(cookie_info)
    
    # ---- 16. PII SCAN ON PAGE TEXT (with false-positive filtering) ----
    # Only flag ACTUAL personal data — not brand names, org names, etc.
    # Organization names, locations, and dates are PUBLIC info, not PII exposure
    REAL_PII_TYPES = {
        "EMAIL_ADDRESS", "PHONE_NUMBER", "US_SSN", "CREDIT_CARD",
        "US_DRIVER_LICENSE", "US_PASSPORT", "US_BANK_NUMBER",
        "IBAN_CODE", "IP_ADDRESS", "MEDICAL_LICENSE",
        "UK_NHS", "SG_NRIC_FIN", "AU_ABN", "AU_ACN",
    }
    
    pii_in_text = []
    text_preview = visible_text[:5000]
    if text_preview.strip() and analyzer:
        try:
            results = analyzer.analyze(
                text=text_preview,
                language="en",
                score_threshold=0.7,  # Higher threshold to reduce false positives
            )
            for r in results:
                # Only flag REAL PII types — skip ORGANIZATION, LOCATION, DATE, PERSON
                # Those are public information on a website, not PII exposure
                if r.entity_type not in REAL_PII_TYPES:
                    continue
                entity_text = text_preview[r.start:r.end].strip()
                if len(entity_text) > 3:
                    pii_in_text.append({
                        "type": r.entity_type,
                        "text": entity_text[:50],
                        "score": round(r.score, 2),
                        "label": ENTITY_META.get(r.entity_type, {}).get("label", r.entity_type),
                    })
            # Deduplicate by text value
            seen_pii = set()
            unique_pii = []
            for p in pii_in_text:
                if p["text"] not in seen_pii:
                    seen_pii.add(p["text"])
                    unique_pii.append(p)
            pii_in_text = unique_pii[:20]
        except Exception as e:
            print(f"[!] PII scan on URL content failed: {e}")
    
    # ---- 9. CALCULATE RISK SCORE ----
    risk_score = 0
    risk_factors = []
    
    if not is_https:
        risk_score += 25
        risk_factors.append("No HTTPS — data transmitted in plain text")
    if len(trackers_found) > 5:
        risk_score += 20
        risk_factors.append(f"{len(trackers_found)} third-party trackers detected")
    elif len(trackers_found) > 0:
        risk_score += 10
        risk_factors.append(f"{len(trackers_found)} third-party tracker(s) found")
    if tracker_categories.get("session_recording", 0) > 0:
        risk_score += 15
        risk_factors.append("Session recording detected — keystrokes/mouse may be captured")
    if tracker_categories.get("fingerprinting", 0) > 0:
        risk_score += 15
        risk_factors.append("Browser fingerprinting detected")
    if len(tracking_pixels) > 0:
        risk_score += 10
        risk_factors.append(f"{len(tracking_pixels)} hidden tracking pixel(s)")
    if len(ai_endpoints_found) > 0:
        key_leaks = [a for a in ai_endpoints_found if a["type"] == "api_key_leak"]
        if key_leaks:
            risk_score += 25
            risk_factors.append(f"Exposed AI API key(s) in client-side code!")
        else:
            risk_score += 5
            risk_factors.append("AI/LLM API endpoints referenced in client code")
    if len(pii_in_text) > 0:
        risk_score += 15
        risk_factors.append(f"{len(pii_in_text)} PII item(s) exposed in page content")
    if not has_privacy_policy:
        risk_score += 10
        risk_factors.append("No privacy policy link found")
    if not has_cookie_consent and len(trackers_found) > 0:
        risk_score += 10
        risk_factors.append("Trackers present but no cookie consent mechanism")
    if len(unique_pii_inputs) > 3:
        risk_score += 5
        risk_factors.append(f"Collects {len(unique_pii_inputs)} types of personal data via forms")
    
    # Blacklight-grade risk factors
    if canvas_fingerprinting:
        risk_score += 15
        risk_factors.append(f"Canvas fingerprinting detected ({len(canvas_fp_signals)} API signals)")
    if key_logging_detected:
        risk_score += 20
        risk_factors.append(f"Key logging detected — keystrokes captured before form submission")
    if session_recording_detected:
        risk_score += 15
        risk_factors.append(f"Session recording — mouse movements/clicks/scrolls being captured")
    if fb_pixel_detected:
        risk_score += 10
        risk_factors.append(f"Facebook Pixel tracking {len(fb_pixel_events)} event type(s)")
    if ga_detected and "user_id" in ga_events:
        risk_score += 10
        risk_factors.append("Google Analytics with user-level tracking (user_id)")
    elif ga_detected:
        risk_score += 5
        risk_factors.append(f"Google Analytics tracking {len(ga_events)} event type(s)")
    if len(third_party_domains_found) > 5:
        risk_score += 10
        risk_factors.append(f"{len(third_party_domains_found)} known ad/tracking domains from Disconnect.me list")
    elif len(third_party_domains_found) > 0:
        risk_score += 5
        risk_factors.append(f"{len(third_party_domains_found)} known ad/tracking domain(s)")
    
    risk_score = min(risk_score, 100)
    
    if risk_score >= 70:
        risk_level = "critical"
    elif risk_score >= 40:
        risk_level = "high"
    elif risk_score >= 20:
        risk_level = "medium"
    else:
        risk_level = "low"
    
    elapsed = round((time.time() - start_time) * 1000, 1)
    
    # ---- BUILD REPORT ----
    report = {
        "url": final_url,
        "domain": base_domain,
        "scanned_at": datetime.now(timezone.utc).isoformat(),
        "scan_time_ms": elapsed,
        "status_code": status_code,
        
        # Risk assessment
        "risk_score": risk_score,
        "risk_level": risk_level,
        "risk_factors": risk_factors,
        
        # Findings
        "trackers": {
            "count": len(trackers_found),
            "items": trackers_found,
            "categories": tracker_categories,
        },
        "tracking_pixels": {
            "count": len(tracking_pixels),
            "items": tracking_pixels[:10],
        },
        "pii_collection": {
            "form_count": len(all_forms),
            "pii_input_count": len(unique_pii_inputs),
            "inputs": unique_pii_inputs,
        },
        "exposed_pii": {
            "count": len(pii_in_text),
            "items": pii_in_text,
        },
        "ai_endpoints": {
            "count": len(ai_endpoints_found),
            "items": ai_endpoints_found,
        },
        
        # Blacklight-grade deep analysis
        "blacklight": {
            "canvas_fingerprinting": {
                "detected": canvas_fingerprinting,
                "signals": canvas_fp_signals[:10],
                "signal_count": len(canvas_fp_signals),
            },
            "key_logging": {
                "detected": key_logging_detected,
                "signals": keylog_signals[:10],
                "signal_count": len(keylog_signals),
            },
            "session_recording": {
                "detected": session_recording_detected,
                "signals": session_rec_signals[:10],
                "signal_count": len(session_rec_signals),
            },
            "facebook_pixel": {
                "detected": fb_pixel_detected,
                "events": fb_pixel_events[:10],
            },
            "google_analytics": {
                "detected": ga_detected,
                "events": ga_events[:10],
                "user_tracking": "user_id" in ga_events,
            },
            "tracking_domains": {
                "count": len(third_party_domains_found),
                "domains": third_party_domains_found[:20],
            },
        },
        
        # Compliance
        "compliance": {
            "https": is_https,
            "privacy_policy": has_privacy_policy,
            "cookie_consent": has_cookie_consent,
            "terms_of_service": has_terms,
            "security_headers": security_headers,
            "security_header_grade": sec_header_grade,
        },
        
        # DPDP Act 2023 (India) Compliance
        "dpdp": {
            "score": dpdp_score,
            "grade": dpdp_grade,
            "total_checks": len(dpdp_checks),
            "checks": {k: {"passed": v, "section": {
                "consent_mechanism": "Section 6 — Consent",
                "privacy_notice": "Section 5 — Notice",
                "grievance_officer": "Section 8(7) — Grievance Redressal",
                "data_retention_policy": "Section 8(6) — Data Retention",
                "children_protection": "Section 9 — Children's Data",
                "consent_withdrawal": "Section 6(4) — Consent Withdrawal",
                "breach_notification": "Section 8(5) — Breach Notification",
            }.get(k, "")} for k, v in dpdp_checks.items()},
        },
        
        # Page info
        "page": {
            "title": (soup_full.title.string.strip() if soup_full.title and soup_full.title.string else ""),
            "text_length": len(visible_text),
            "scripts_count": len(all_scripts),
            "forms_count": len(all_forms),
            "images_count": len(all_imgs),
        },
        
        # Scan engine info
        "engine": {
            "text_extraction": "Jina Reader API (JS-rendered)" if jina_used else "BeautifulSoup (static HTML)",
            "html_analysis": "requests + BeautifulSoup",
            "pii_detection": "Microsoft Presidio NLP" if analyzer else "unavailable",
            "methodology": "Blacklight (The Markup) + DPDP Act 2023 + Presidio + Jina Reader",
        },
        
        # Cookie deep analysis
        "cookies": {
            "count": len(cookie_analysis),
            "items": cookie_analysis[:20],
            "summary": {
                "session_cookies": sum(1 for c in cookie_analysis if not c.get("persistent")),
                "persistent_cookies": sum(1 for c in cookie_analysis if c.get("persistent")),
                "httponly": sum(1 for c in cookie_analysis if c.get("httponly")),
                "secure": sum(1 for c in cookie_analysis if c.get("secure")),
                "samesite": sum(1 for c in cookie_analysis if c.get("samesite")),
                "third_party": sum(1 for c in cookie_analysis if c.get("third_party")),
            },
        },
    }
    
    return report


# ---- Serve Frontend Static Files ----
# Serve static files from the current directory
app.mount("/static", StaticFiles(directory="."), name="static")

@app.get("/")
def serve_index():
    return FileResponse("index.html")

@app.get("/dashboard")
@app.get("/dashboard.html")
def serve_dashboard():
    return FileResponse("dashboard.html")

# Catch-all for CSS/JS files
@app.get("/{filename}")
def serve_file(filename: str):
    filepath = os.path.join(".", filename)
    if os.path.isfile(filepath):
        return FileResponse(filepath)
    raise HTTPException(404, "Not found")


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    print(f"\n[*] RedactAI API Server starting on port {port}...")
    print(f"[>] Dashboard: http://127.0.0.1:{port}/dashboard")
    print(f"[>] API Docs:  http://127.0.0.1:{port}/docs")
    print(f"[>] Landing:   http://127.0.0.1:{port}/\n")
    uvicorn.run(app, host="0.0.0.0", port=port)
