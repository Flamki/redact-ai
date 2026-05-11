"""
RedactAI — FastAPI Backend powered by Microsoft Presidio
Production-grade PII detection & redaction API
"""

import os
import json
import time
import uuid
import hashlib
from datetime import datetime, timezone
from typing import Optional
from fastapi import FastAPI, File, UploadFile, HTTPException, Header, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel

# ---- Supabase Setup ----
SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://amklcfiyxeomdueeptyu.supabase.co")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImFta2xjZml5eGVvbWR1ZWVwdHl1Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3Nzg1MDQ2OTIsImV4cCI6MjA5NDA4MDY5Mn0.WNvPc9hrorOw_pMI2PS8pVPklfqwXCQH3kBJSwja6dk")
SUPABASE_AVAILABLE = False
try:
    from supabase import create_client
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    # Quick connectivity check
    supabase.table("redact_scans").select("id").limit(1).execute()
    SUPABASE_AVAILABLE = True
    print("[+] Supabase connected! Persistent history enabled.")
except Exception as e:
    print(f"[!] Supabase unavailable ({e}), falling back to in-memory history")
    supabase = None

# ---- Presidio Setup ----
from presidio_analyzer import AnalyzerEngine, RecognizerRegistry
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

# ---- Piiranha Transformer Recognizer ----
from presidio_analyzer import EntityRecognizer, RecognizerResult

PIIRANHA_AVAILABLE = False
try:
    from transformers import pipeline as hf_pipeline
    PIIRANHA_AVAILABLE = True
except ImportError:
    print("[!] transformers not installed, skipping Piiranha model")

class PiiranhaRecognizer(EntityRecognizer):
    """Custom Presidio recognizer using the Piiranha PII model (DeBERTa-v3, 99.4% accuracy)"""

    PIIRANHA_TO_PRESIDIO = {
        "GIVENNAME": "PERSON",
        "SURNAME": "PERSON",
        "FIRSTNAME": "PERSON",
        "LASTNAME": "PERSON",
        "EMAIL": "EMAIL_ADDRESS",
        "PHONE": "PHONE_NUMBER",
        "PHONENUMBER": "PHONE_NUMBER",
        "CREDITCARD": "CREDIT_CARD",
        "CREDITCARDNUMBER": "CREDIT_CARD",
        "SOCIALNUM": "US_SSN",
        "SOCIALSECURITYNUMBER": "US_SSN",
        "DRIVERSLICENSE": "US_DRIVER_LICENSE",
        "DATEOFBIRTH": "DATE_TIME",
        "DOB": "DATE_TIME",
        "IDCARD": "ID_CARD",
        "TAXNUMBER": "TAX_ID",
        "STREETADDRESS": "LOCATION",
        "CITY": "LOCATION",
        "ZIPCODE": "LOCATION",
        "BUILDINGNUMBER": "LOCATION",
        "ACCOUNTNUMBER": "ACCOUNT_NUMBER",
        "USERNAME": "USERNAME",
        "PASSWORD": "PASSWORD",
    }

    def __init__(self):
        supported = list(set(self.PIIRANHA_TO_PRESIDIO.values()))
        super().__init__(
            supported_entities=supported,
            supported_language="en",
            name="PiiranhaRecognizer",
        )
        print("[*] Loading Piiranha PII transformer model...")
        self.pipe = hf_pipeline(
            "token-classification",
            model="iiiorg/piiranha-v1-detect-personal-information",
            aggregation_strategy="max",
            device=-1,  # CPU
        )
        print("[+] Piiranha model loaded!")

    def load(self):
        pass

    def analyze(self, text, entities=None, nlp_artifacts=None):
        results = []
        try:
            preds = self.pipe(text)
            for pred in preds:
                label = pred["entity_group"].upper().replace("-", "")
                presidio_type = self.PIIRANHA_TO_PRESIDIO.get(label, None)
                if presidio_type and (entities is None or presidio_type in entities):
                    results.append(
                        RecognizerResult(
                            entity_type=presidio_type,
                            start=pred["start"],
                            end=pred["end"],
                            score=round(float(pred["score"]), 3),
                        )
                    )
        except Exception as e:
            print(f"[!] Piiranha error: {e}")
        return results


# Initialize engines — try large model, fall back to small
print("[*] Loading NLP model & Presidio engines...")

# Try en_core_web_lg first (better NER), fall back to en_core_web_sm
for model_name in ["en_core_web_lg", "en_core_web_sm"]:
    try:
        nlp_config = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": model_name}],
        }
        nlp_engine = NlpEngineProvider(nlp_configuration=nlp_config).create_engine()
        print(f"[+] Using spaCy model: {model_name}")
        break
    except Exception as e:
        print(f"[!] {model_name} not available: {e}")
        continue

registry = RecognizerRegistry()
registry.load_predefined_recognizers(nlp_engine=nlp_engine)

# Add Piiranha transformer if available
if PIIRANHA_AVAILABLE:
    try:
        piiranha = PiiranhaRecognizer()
        registry.add_recognizer(piiranha)
        print("[+] Piiranha transformer recognizer added!")
    except Exception as e:
        print(f"[!] Could not load Piiranha model: {e}")
        print("[*] Continuing with spaCy-only detection")

# ---- GLiNER Zero-Shot NER (contextual understanding) ----
GLINER_AVAILABLE = False
try:
    from gliner import GLiNER as GLiNERModel
    GLINER_AVAILABLE = True
except ImportError:
    print("[!] gliner not installed, skipping zero-shot NER")

class GLiNERRecognizer(EntityRecognizer):
    """Zero-shot NER using GLiNER — understands context, no training needed.
    Detects entities based on natural language labels like 'person name', 'date'."""

    # Map GLiNER labels → Presidio entity types
    LABEL_MAP = {
        "person name": "PERSON",
        "full name": "PERSON",
        "date": "DATE_TIME",
        "monetary amount": "MONETARY_VALUE",
        "organization": "ORGANIZATION",
        "address": "LOCATION",
        "city": "LOCATION",
        "country": "LOCATION",
    }

    # Entity labels we ask GLiNER to find — written in natural language
    DETECT_LABELS = [
        "person name",
        "date",
        "monetary amount",
        "organization",
        "address",
    ]

    def __init__(self):
        supported = list(set(self.LABEL_MAP.values()))
        super().__init__(
            supported_entities=supported,
            supported_language="en",
            name="GLiNERRecognizer",
        )
        print("[*] Loading GLiNER zero-shot NER model...")
        self.model = GLiNERModel.from_pretrained("urchade/gliner_medium-v2.1")
        print("[+] GLiNER model loaded!")

    def load(self):
        pass

    def analyze(self, text, entities=None, nlp_artifacts=None):
        results = []
        try:
            preds = self.model.predict_entities(text, self.DETECT_LABELS, threshold=0.4)
            for pred in preds:
                label = pred["label"].lower()
                presidio_type = self.LABEL_MAP.get(label, None)
                if presidio_type and (entities is None or presidio_type in entities):
                    results.append(
                        RecognizerResult(
                            entity_type=presidio_type,
                            start=pred["start"],
                            end=pred["end"],
                            score=round(float(pred["score"]), 3),
                        )
                    )
        except Exception as e:
            print(f"[!] GLiNER error: {e}")
        return results

if GLINER_AVAILABLE:
    try:
        gliner_rec = GLiNERRecognizer()
        registry.add_recognizer(gliner_rec)
        print("[+] GLiNER zero-shot recognizer added!")
    except Exception as e:
        print(f"[!] Could not load GLiNER: {e}")

# ---- Custom Informal Date Recognizer ----
import re
from presidio_analyzer import Pattern, PatternRecognizer

MONTHS = r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"

informal_date_patterns = [
    # "9th march", "10th march", "1st january", "23rd april"
    Pattern("ordinal_month", rf"\b\d{{1,2}}(?:st|nd|rd|th)\s+{MONTHS}\b", 0.85),
    # "march 9th", "april 10th", "january 1st"
    Pattern("month_ordinal", rf"\b{MONTHS}\s+\d{{1,2}}(?:st|nd|rd|th)?\b", 0.85),
    # "march 2024", "april 2025"
    Pattern("month_year", rf"\b{MONTHS}\s+\d{{4}}\b", 0.80),
    # "9th march 2024"
    Pattern("ordinal_month_year", rf"\b\d{{1,2}}(?:st|nd|rd|th)\s+{MONTHS}\s+\d{{4}}\b", 0.90),
    # standalone months in context: "in march", "on april", "by december"
    Pattern("standalone_month", rf"\b(?:in|on|by|before|after|since|until|during)\s+{MONTHS}\b", 0.70),
]

date_recognizer = PatternRecognizer(
    supported_entity="DATE_TIME",
    name="InformalDateRecognizer",
    patterns=informal_date_patterns,
    supported_language="en",
)
registry.add_recognizer(date_recognizer)
print("[+] Informal date recognizer added!")

analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry, supported_languages=["en"])
anonymizer = AnonymizerEngine()
print("[+] Presidio engines ready!")

# ---- FastAPI App ----
app = FastAPI(
    title="RedactAI API",
    description="AI-powered PII detection & redaction API backed by Microsoft Presidio",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---- Storage (Supabase persistent + in-memory fallback) ----
scan_history_mem = []  # fallback only
api_keys = {
    "rda_live_sk_demo123": {"name": "Demo Key", "created": datetime.now().isoformat(), "active": True}
}

def save_scan(record):
    """Save a scan record to Supabase (or in-memory fallback)"""
    if SUPABASE_AVAILABLE:
        try:
            supabase.table("redact_scans").insert({
                "source": record["source"],
                "entity_count": record["entity_count"],
                "types": json.dumps(record["types"]),
                "processing_ms": int(record["processing_ms"]),
                "preview": record.get("preview", ""),
            }).execute()
            return
        except Exception as e:
            print(f"[!] Supabase insert failed: {e}")
    scan_history_mem.append(record)

# ---- Entity color/icon mapping for frontend ----
ENTITY_META = {
    "PERSON": {"icon": "👤", "color": "#f472b6", "cssClass": "name", "label": "Person Name"},
    "EMAIL_ADDRESS": {"icon": "📧", "color": "#74c0fc", "cssClass": "email", "label": "Email"},
    "PHONE_NUMBER": {"icon": "📱", "color": "#51cf66", "cssClass": "phone", "label": "Phone"},
    "CREDIT_CARD": {"icon": "💳", "color": "#ffd43b", "cssClass": "credit-card", "label": "Credit Card"},
    "US_SSN": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "SSN"},
    "US_PASSPORT": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Passport"},
    "US_DRIVER_LICENSE": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Driver License"},
    "IP_ADDRESS": {"icon": "🌐", "color": "#22d3ee", "cssClass": "ip", "label": "IP Address"},
    "DATE_TIME": {"icon": "📅", "color": "#a29bfe", "cssClass": "date", "label": "Date/Time"},
    "LOCATION": {"icon": "📍", "color": "#fdcb6e", "cssClass": "location", "label": "Location"},
    "NRP": {"icon": "🏛️", "color": "#dfe6e9", "cssClass": "other", "label": "Nationality/Religion"},
    "MEDICAL_LICENSE": {"icon": "🏥", "color": "#e17055", "cssClass": "other", "label": "Medical License"},
    "URL": {"icon": "🔗", "color": "#74c0fc", "cssClass": "other", "label": "URL"},
    "IBAN_CODE": {"icon": "🏦", "color": "#ffd43b", "cssClass": "credit-card", "label": "IBAN"},
    "CRYPTO": {"icon": "₿", "color": "#f9ca24", "cssClass": "other", "label": "Crypto Wallet"},
    "UK_NHS": {"icon": "🏥", "color": "#e17055", "cssClass": "gov-id", "label": "UK NHS Number"},
    "IN_AADHAAR": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Aadhaar"},
    "IN_PAN": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "PAN Card"},
    "ID_CARD": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "ID Card"},
    "TAX_ID": {"icon": "🆔", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Tax Number"},
    "ACCOUNT_NUMBER": {"icon": "🏦", "color": "#ffd43b", "cssClass": "credit-card", "label": "Account Number"},
    "USERNAME": {"icon": "👤", "color": "#a29bfe", "cssClass": "name", "label": "Username"},
    "PASSWORD": {"icon": "🔒", "color": "#ff6b6b", "cssClass": "gov-id", "label": "Password"},
    "MONETARY_VALUE": {"icon": "💰", "color": "#ffd43b", "cssClass": "credit-card", "label": "Money/Amount"},
    "ORGANIZATION": {"icon": "🏢", "color": "#dfe6e9", "cssClass": "other", "label": "Organization"},
}

# ---- File Text Extraction ----
def extract_text_from_file(content: bytes, ext: str) -> str:
    """Extract text from various file formats"""
    import io
    
    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            print(f"[!] PDF extraction failed: {e}")
            return ""
    
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"[!] DOCX extraction failed: {e}")
            return ""
    
    elif ext in ("xlsx", "xls"):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            text_parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        text_parts.append(" ".join(cells))
            wb.close()
            return "\n".join(text_parts)
        except Exception as e:
            print(f"[!] XLSX extraction failed: {e}")
            return ""
    
    elif ext == "csv":
        import csv
        text = content.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        return " ".join(" ".join(row) for row in reader)
    
    elif ext == "json":
        text = content.decode("utf-8", errors="ignore")
        try:
            data = json.loads(text)
            return json.dumps(data) if isinstance(data, (dict, list)) else text
        except json.JSONDecodeError:
            return text
    
    else:  # txt and fallback
        return content.decode("utf-8", errors="ignore")

# ---- Request/Response Models ----
class ScanRequest(BaseModel):
    text: str
    mode: str = "highlight"  # "highlight" or "redact"
    language: str = "en"
    entities: Optional[list] = None  # specific entities to detect, or None for all
    score_threshold: float = 0.35

class ScanResponse(BaseModel):
    original: str
    redacted: str
    entities: list
    entity_summary: dict
    count: int
    processing_ms: float

class BatchScanRequest(BaseModel):
    texts: list[str]
    mode: str = "redact"
    language: str = "en"

# ---- API Endpoints ----

@app.get("/api/health")
def health_check():
    return {"status": "healthy", "engine": "presidio", "version": "1.0.0"}


@app.post("/api/v1/scan", response_model=ScanResponse)
def scan_text(req: ScanRequest):
    """Scan text for PII and return detected entities + redacted text"""
    start = time.time()
    
    # Analyze with Presidio
    results = analyzer.analyze(
        text=req.text,
        language=req.language,
        entities=req.entities,
        score_threshold=req.score_threshold,
    )
    
    # Build entity list with metadata
    entities = []
    for r in sorted(results, key=lambda x: x.start):
        meta = ENTITY_META.get(r.entity_type, {"icon": "❓", "color": "#dfe6e9", "cssClass": "other", "label": r.entity_type})
        entities.append({
            "type": r.entity_type,
            "label": meta["label"],
            "text": req.text[r.start:r.end],
            "start": r.start,
            "end": r.end,
            "score": round(float(r.score), 3),
            "icon": meta["icon"],
            "color": meta["color"],
            "cssClass": meta["cssClass"],
        })
    
    # Anonymize/redact
    anonymized = anonymizer.anonymize(
        text=req.text,
        analyzer_results=results,
        operators={
            "DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"}),
            "PERSON": OperatorConfig("replace", {"new_value": "[NAME]"}),
            "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[EMAIL]"}),
            "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[PHONE]"}),
            "CREDIT_CARD": OperatorConfig("replace", {"new_value": "[CREDIT_CARD]"}),
            "US_SSN": OperatorConfig("replace", {"new_value": "[SSN]"}),
            "IP_ADDRESS": OperatorConfig("replace", {"new_value": "[IP_ADDRESS]"}),
            "DATE_TIME": OperatorConfig("replace", {"new_value": "[DATE]"}),
            "LOCATION": OperatorConfig("replace", {"new_value": "[LOCATION]"}),
            "URL": OperatorConfig("replace", {"new_value": "[URL]"}),
            "IN_AADHAAR": OperatorConfig("replace", {"new_value": "[AADHAAR]"}),
            "IN_PAN": OperatorConfig("replace", {"new_value": "[PAN]"}),
        }
    )
    
    # Build entity summary
    summary = {}
    for e in entities:
        t = e["label"]
        if t not in summary:
            summary[t] = {"count": 0, "icon": e["icon"], "cssClass": e["cssClass"]}
        summary[t]["count"] += 1
    
    elapsed_ms = round((time.time() - start) * 1000, 2)
    
    # Store in history (persistent via Supabase)
    save_scan({
        "id": str(uuid.uuid4())[:8],
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "source": "Text Input",
        "entity_count": len(entities),
        "types": list(summary.keys()),
        "processing_ms": elapsed_ms,
        "preview": req.text[:80] + ("..." if len(req.text) > 80 else ""),
    })
    
    return ScanResponse(
        original=req.text,
        redacted=anonymized.text,
        entities=entities,
        entity_summary=summary,
        count=len(entities),
        processing_ms=elapsed_ms,
    )


@app.post("/api/v1/scan/batch")
def scan_batch(req: BatchScanRequest):
    """Scan multiple texts at once"""
    results = []
    total_start = time.time()
    
    for text in req.texts:
        analysis = analyzer.analyze(text=text, language=req.language)
        anonymized = anonymizer.anonymize(text=text, analyzer_results=analysis)
        
        entities = []
        for r in analysis:
            meta = ENTITY_META.get(r.entity_type, {"icon": "❓", "label": r.entity_type})
            entities.append({
                "type": r.entity_type,
                "label": meta["label"],
                "text": text[r.start:r.end],
                "score": round(float(r.score), 3),
            })
        
        results.append({
            "original": text,
            "redacted": anonymized.text,
            "entity_count": len(entities),
            "entities": entities,
        })
    
    return {
        "results": results,
        "total_texts": len(req.texts),
        "total_entities": sum(r["entity_count"] for r in results),
        "processing_ms": round((time.time() - total_start) * 1000, 2),
    }


@app.post("/api/v1/scan/file")
async def scan_file(file: UploadFile = File(...)):
    """Upload and scan a file for PII — supports TXT, CSV, JSON, PDF, DOCX, XLSX"""
    if not file.filename:
        raise HTTPException(400, "No file provided")
    
    ext = file.filename.rsplit(".", 1)[-1].lower()
    supported = ("txt", "csv", "json", "pdf", "docx", "doc", "xlsx", "xls")
    if ext not in supported:
        raise HTTPException(400, f"Unsupported file type: .{ext}. Supported: {', '.join(supported)}")
    
    content = await file.read()
    start = time.time()
    
    # Extract text based on file type
    all_text = extract_text_from_file(content, ext)
    if not all_text or not all_text.strip():
        raise HTTPException(400, "Could not extract text from file")
    
    # Analyze
    results = analyzer.analyze(text=all_text, language="en")
    anonymized = anonymizer.anonymize(text=all_text, analyzer_results=results)
    
    entities = []
    for r in sorted(results, key=lambda x: x.start):
        meta = ENTITY_META.get(r.entity_type, {"icon": "❓", "label": r.entity_type})
        entities.append({
            "type": r.entity_type,
            "label": meta["label"],
            "text": all_text[r.start:r.end],
            "score": round(float(r.score), 3),
        })
    
    elapsed_ms = round((time.time() - start) * 1000, 2)
    
    # Store in history (persistent via Supabase)
    save_scan({
        "id": str(uuid.uuid4())[:8],
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "source": f"File: {file.filename}",
        "entity_count": len(entities),
        "types": list(set(e["label"] for e in entities)),
        "processing_ms": elapsed_ms,
        "preview": all_text[:80] + "...",
    })
    
    return {
        "filename": file.filename,
        "file_size": len(content),
        "redacted_text": anonymized.text,
        "entities": entities,
        "entity_count": len(entities),
        "processing_ms": elapsed_ms,
    }


@app.get("/api/v1/history")
def get_history(page: int = 1, per_page: int = 10):
    """Get scan history with pagination — reads from Supabase"""
    if SUPABASE_AVAILABLE:
        try:
            count_resp = supabase.table("redact_scans").select("id", count="exact").execute()
            total = count_resp.count or 0
            offset = (page - 1) * per_page
            data_resp = supabase.table("redact_scans") \
                .select("*") \
                .order("created_at", desc=True) \
                .range(offset, offset + per_page - 1) \
                .execute()
            items = []
            for row in data_resp.data:
                types_val = row.get("types", "[]")
                if isinstance(types_val, str):
                    try:
                        types_val = json.loads(types_val)
                    except Exception:
                        types_val = []
                items.append({
                    "id": str(row["id"])[:8],
                    "timestamp": row["created_at"],
                    "source": row.get("source", "Unknown"),
                    "entity_count": row.get("entity_count", 0),
                    "types": types_val,
                    "processing_ms": row.get("processing_ms", 0),
                    "preview": row.get("preview", ""),
                })
            return {
                "items": items,
                "total": total,
                "page": page,
                "pages": max(1, (total + per_page - 1) // per_page),
            }
        except Exception as e:
            print(f"[!] Supabase history read failed: {e}")
    # Fallback to in-memory
    total = len(scan_history_mem)
    start = (page - 1) * per_page
    items = list(reversed(scan_history_mem))[start:start + per_page]
    return {
        "items": items,
        "total": total,
        "page": page,
        "pages": max(1, (total + per_page - 1) // per_page),
    }


@app.get("/api/v1/stats")
def get_stats():
    """Get overview statistics — reads from Supabase"""
    if SUPABASE_AVAILABLE:
        try:
            count_resp = supabase.table("redact_scans").select("id", count="exact").execute()
            total_scans = count_resp.count or 0
            all_resp = supabase.table("redact_scans").select("entity_count,processing_ms,types").execute()
            rows = all_resp.data or []
            total_entities = sum(r.get("entity_count", 0) for r in rows)
            avg_ms = round(sum(r.get("processing_ms", 0) for r in rows) / max(1, total_scans), 2)
            type_counts = {}
            for r in rows:
                types_val = r.get("types", "[]")
                if isinstance(types_val, str):
                    try:
                        types_val = json.loads(types_val)
                    except Exception:
                        types_val = []
                for t in types_val:
                    type_counts[t] = type_counts.get(t, 0) + 1
            return {
                "total_scans": total_scans,
                "total_entities": total_entities,
                "avg_response_ms": avg_ms,
                "entity_type_breakdown": type_counts,
            }
        except Exception as e:
            print(f"[!] Supabase stats read failed: {e}")
    # Fallback to in-memory
    total_scans = len(scan_history_mem)
    total_entities = sum(h["entity_count"] for h in scan_history_mem)
    avg_ms = round(sum(h["processing_ms"] for h in scan_history_mem) / max(1, total_scans), 2)
    type_counts = {}
    for h in scan_history_mem:
        for t in h.get("types", []):
            type_counts[t] = type_counts.get(t, 0) + 1
    return {
        "total_scans": total_scans,
        "total_entities": total_entities,
        "avg_response_ms": avg_ms,
        "entity_type_breakdown": type_counts,
    }


@app.get("/api/v1/supported-entities")
def get_supported_entities():
    """List all PII entity types the engine can detect"""
    supported = analyzer.get_supported_entities()
    entities = []
    for entity_type in sorted(supported):
        meta = ENTITY_META.get(entity_type, {"icon": "❓", "color": "#dfe6e9", "label": entity_type})
        entities.append({
            "type": entity_type,
            "label": meta["label"],
            "icon": meta["icon"],
            "color": meta["color"],
        })
    return {"entities": entities, "count": len(entities)}


class CustomDetectorRequest(BaseModel):
    name: str
    entity_type: str
    regex: str
    score: float = 0.8


@app.post("/api/v1/custom-detector")
def add_custom_detector(req: CustomDetectorRequest):
    """Register a custom regex-based PII detector at runtime"""
    import re
    
    # Validate regex
    try:
        re.compile(req.regex)
    except re.error as e:
        raise HTTPException(400, f"Invalid regex: {e}")
    
    # Use Presidio's PatternRecognizer API
    from presidio_analyzer import Pattern, PatternRecognizer
    
    pattern = Pattern(
        name=req.name,
        regex=req.regex,
        score=req.score,
    )
    
    recognizer = PatternRecognizer(
        supported_entity=req.entity_type,
        name=f"custom_{req.name.lower().replace(' ', '_')}",
        patterns=[pattern],
    )
    
    # Add to the live registry
    analyzer.registry.add_recognizer(recognizer)
    
    # Also add to entity meta for frontend display
    ENTITY_META[req.entity_type] = {
        "icon": "🔧",
        "color": "#b8e994",
        "cssClass": "other",
        "label": req.name,
    }
    
    return {
        "status": "ok",
        "message": f"Custom detector '{req.name}' registered for entity '{req.entity_type}'",
        "entity_type": req.entity_type,
        "pattern": req.regex,
    }


@app.get("/api/v1/custom-detectors")
def list_custom_detectors():
    """List all custom detectors currently registered"""
    custom = []
    for rec in analyzer.registry.recognizers:
        if hasattr(rec, 'name') and rec.name and rec.name.startswith('custom_'):
            patterns = []
            if hasattr(rec, 'patterns'):
                patterns = [{"name": p.name, "regex": p.regex, "score": p.score} for p in rec.patterns]
            custom.append({
                "name": rec.name,
                "entity_type": rec.supported_entities[0] if rec.supported_entities else "UNKNOWN",
                "patterns": patterns,
            })
    return {"detectors": custom, "count": len(custom)}


@app.get("/api/v1/export")
def export_history(format: str = "csv"):
    """Export scan history as CSV or JSON — for compliance/audit"""
    import io
    
    # Fetch all history
    items = []
    if SUPABASE_AVAILABLE:
        try:
            resp = supabase.table("redact_scans") \
                .select("*") \
                .order("created_at", desc=True) \
                .limit(1000) \
                .execute()
            for row in resp.data:
                types_val = row.get("types", "[]")
                if isinstance(types_val, str):
                    try:
                        types_val = json.loads(types_val)
                    except Exception:
                        types_val = []
                items.append({
                    "id": str(row["id"])[:8],
                    "timestamp": row["created_at"],
                    "source": row.get("source", ""),
                    "entity_count": row.get("entity_count", 0),
                    "types": ", ".join(types_val) if types_val else "",
                    "processing_ms": row.get("processing_ms", 0),
                    "preview": row.get("preview", ""),
                })
        except Exception as e:
            print(f"[!] Export from Supabase failed: {e}")
    else:
        for h in reversed(scan_history_mem):
            items.append({
                "id": h.get("id", ""),
                "timestamp": h.get("timestamp", ""),
                "source": h.get("source", ""),
                "entity_count": h.get("entity_count", 0),
                "types": ", ".join(h.get("types", [])),
                "processing_ms": h.get("processing_ms", 0),
                "preview": h.get("preview", ""),
            })
    
    if format == "json":
        return JSONResponse(content={"export": items, "total": len(items), "exported_at": datetime.now(timezone.utc).isoformat()})
    
    # CSV format
    import csv
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=["id", "timestamp", "source", "entity_count", "types", "processing_ms", "preview"])
    writer.writeheader()
    writer.writerows(items)
    
    from fastapi.responses import StreamingResponse
    csv_content = output.getvalue()
    return StreamingResponse(
        iter([csv_content]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=redactai_audit_log_{datetime.now().strftime('%Y%m%d')}.csv"}
    )


# ---- Serve Frontend Static Files ----
# Serve static files from the current directory
app.mount("/static", StaticFiles(directory="."), name="static")

@app.get("/")
def serve_index():
    return FileResponse("index.html")

@app.get("/dashboard")
@app.get("/dashboard.html")
def serve_dashboard():
    return FileResponse("dashboard.html")

# Catch-all for CSS/JS files
@app.get("/{filename}")
def serve_file(filename: str):
    filepath = os.path.join(".", filename)
    if os.path.isfile(filepath):
        return FileResponse(filepath)
    raise HTTPException(404, "Not found")


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    print(f"\n[*] RedactAI API Server starting on port {port}...")
    print(f"[>] Dashboard: http://127.0.0.1:{port}/dashboard")
    print(f"[>] API Docs:  http://127.0.0.1:{port}/docs")
    print(f"[>] Landing:   http://127.0.0.1:{port}/\n")
    uvicorn.run(app, host="0.0.0.0", port=port)
